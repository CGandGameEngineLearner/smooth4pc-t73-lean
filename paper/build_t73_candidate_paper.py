from __future__ import annotations

from pathlib import Path
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


REPO = Path(r"D:\toffee_code_in_Cursor\smooth4pc-t73-lean\.worktrees\formalize-t73")
OUT = REPO / "paper"
PDF = OUT / "T73_SPC4_CANDIDATE_FALSIFICATION_20260902.pdf"
DOCX = OUT / "T73_SPC4_CANDIDATE_FALSIFICATION_20260902.docx"


TITLE = "A Candidate Degree-494 Skein-Lasagna Obstruction for the Cappell–Shaneson Sphere X(41,189,73)"
AUTHOR = ""
DATE = ""


def para(text: str) -> dict:
    return {"kind": "para", "text": text}


def section(title: str) -> dict:
    return {"kind": "section", "text": title}


def subsection(title: str) -> dict:
    return {"kind": "subsection", "text": title}


def equation(text: str, number: str | None = None) -> dict:
    return {"kind": "equation", "text": text, "number": number}


def theorem(label: str, text: str) -> dict:
    return {"kind": "theorem", "label": label, "text": text}


def bullets(items: list[str]) -> dict:
    return {"kind": "bullets", "items": items}


def code(lines: list[str]) -> dict:
    return {"kind": "code", "lines": lines}


def table(headers: list[str], rows: list[list[str]], widths: list[float]) -> dict:
    return {"kind": "table", "headers": headers, "rows": rows, "widths": widths}


def pretty_math(text: str) -> str:
    replacements = (
        ("S^2_{0,494}", "S²₀,₄₉₄"), ("S^2_0", "S²₀"), ("S^4", "S⁴"),
        ("Lambda^2", "Λ²"),
        ("D_3", "D₃"), ("HH_0", "HH₀"), ("P_86", "P₈₆"),
        ("F_84", "F₈₄"), ("h^3", "h³"), ("h^4", "h⁴"),
        ("h^2", "h²"), ("X^227", "X²²⁷"), ("X^2", "X²"),
        ("t=zeta^{-2}", "t=ζ⁻²"), ("C_s", "Cₛ"), ("s_0", "s₀"),
        ("m_2", "m₂"), ("T_1", "T₁"), ("T_0", "T₀"),
        ("B_act^{-1}", "B_act⁻¹"), ("eta_R", "η_R"),
        ("D_j", "Dⱼ"), ("U_j", "Uⱼ"), ("E_j", "Eⱼ"), ("b_j", "bⱼ"),
        ("e_0-e_5", "e₀−e₅"), ("R_2h", "R₂h"), ("R_3h", "R₃h"),
        ("ell0", "ℓ₀"), ("ell1", "ℓ₁"), ("ell2", "ℓ₂"),
        ("q01", "q₀₁"), ("q12", "q₁₂"), ("rho_h", "ρ_h"),
        ("Theta_h", "Θ_h"), ("A-I", "A−I"), ("W-I", "W−I"),
        ("t-1", "t−1"), ("q-1", "q−1"), ("-59072", "−59072"),
        ("Cappell-Shaneson", "Cappell–Shaneson"),
        ("Khovanov-Rozansky", "Khovanov–Rozansky"),
        ("Horvat-Jablonowski", "Horvat–Jabłonowski"),
        ("Manolescu-Neithalath", "Manolescu–Neithalath"),
        ("Poincare", "Poincaré"),
        ("!=", "≠"), (" -> ", " → "), (" subset ", " ⊂ "),
    )
    for old, new in replacements:
        text = text.replace(old, new)
    for old, new in (("epsilon", "ε"), ("rho", "ρ"), ("zeta", "ζ"),
                     ("Lambda", "Λ"), ("Psi", "Ψ"), ("Sigma", "Σ"),
                     ("Theta", "Θ"), ("alpha", "α")):
        text = re.sub(rf"\b{old}\b", new, text)
    text = re.sub(r"(?<![A-Za-z0-9_])-(?=\d)", "−", text)
    text = re.sub(r"(?<=[0-9)])-(?=\d)", "−", text)
    text = re.sub(r"(?<=[²³⁴⁵⁶⁷⁸⁹])-(?=\d)", "−", text)
    text = text.replace(" - ", " − ")
    return text


SUPER_CHARS = {
    "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
    "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
    "⁻": "−", "⁽": "(", "⁾": ")", "ʳ": "r", "ᴿ": "R",
}
SUB_CHARS = {
    "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
    "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
    "ₛ": "s", "ₜ": "t", "ₑ": "e", "ⱼ": "j", "ₐ": "a",
}
MATH_SYMBOLS = "⊗⊕∈⊂≅∧∪∅↦𝒮ℓΛΨΣΘΓξζρεηπχ"


def scripts_to_markup(text: str) -> str:
    out: list[str] = []
    index = 0
    while index < len(text):
        char = text[index]
        if char in SUPER_CHARS:
            run = []
            while index < len(text) and text[index] in SUPER_CHARS:
                run.append(SUPER_CHARS[text[index]])
                index += 1
            out.append("<super>" + "".join(run) + "</super>")
            continue
        if char in SUB_CHARS:
            run = []
            while index < len(text) and text[index] in SUB_CHARS:
                run.append(SUB_CHARS[text[index]])
                index += 1
            out.append("<sub>" + "".join(run) + "</sub>")
            continue
        out.append(char)
        index += 1
    return "".join(out)


def pdf_markup(text: str) -> str:
    rendered = scripts_to_markup(escape(pretty_math(text)))
    rendered = rendered.replace("v_T", "v<sub>T</sub>")
    rendered = rendered.replace("^quot", "<super>quot</super>")
    rendered = re.sub(r"\^\{([^{}]+)\}", r"<super>\1</super>", rendered)
    rendered = re.sub(r"_\{([^{}]+)\}", r"<sub>\1</sub>", rendered)
    return re.sub(
        "([" + re.escape(MATH_SYMBOLS) + "])",
        r'<font name="CambriaMath">\1</font>',
        rendered,
    )


def code_markup(text: str) -> str:
    rendered = escape(text)
    url = "https://github.com/toffee-desuwa/smooth4pc-t73-lean.git"
    if url in rendered:
        rendered = rendered.replace(
            url,
            f'<a href="{url}" color="#1A4B8C">{url}</a>',
        )
    return rendered


CONTENT: list[dict] = [
    {"kind": "abstract", "text": (
        "We present a candidate obstruction to the smooth four-dimensional Poincare conjecture for the "
        "Cappell-Shaneson manifold X(41,189,73). The proposed distinguishing class lies in homological "
        "degree zero and quantum degree 494 of the N=2 skein-lasagna module over Q. A divided cubic "
        "functional evaluates on the selected class as -59072, while the standard four-sphere module "
        "vanishes in that quantum degree. The argument constructs a balanced "
        "one-handle class, descends its detector through the full beta/psi relations and six relations "
        "from a chosen three-sphere basis, and then invokes the published four-handle and standard-sphere "
        "comparisons. Exact finite calculations, immutable geometry certificates, and a Lean kernel audit "
        "are supplied. The Lean development checks the arithmetic and the abstract quotient implication; "
        "candidate-specific geometric identifications remain explicit interfaces and therefore require "
        "independent human verification. This draft is not peer reviewed."
    )},
    {"kind": "keywords", "text": (
        "Smooth four-dimensional Poincare conjecture; Cappell-Shaneson sphere; skein-lasagna module; "
        "Khovanov-Rozansky homology; handle decomposition; formal verification."
    )},

    section("1. Introduction"),
    para(
        "The smooth four-dimensional Poincare conjecture asks whether every smooth, closed four-manifold "
        "homotopy equivalent to S^4 is diffeomorphic to S^4. The Cappell-Shaneson construction produces a "
        "large and unusually explicit family of homotopy four-spheres from matrices in SL(3,Z). Many members "
        "of the family have been shown standard by Kirby-calculus arguments, but the presentation remains a "
        "natural testing ground for invariants that can distinguish smooth structures."
    ),
    para(
        "This paper studies the member X(41,189,73). The proposed obstruction is not a conventional numerical "
        "invariant attached directly to the matrix. It is a nonzero class in a graded skein-lasagna module, "
        "detected only after a coefficient trace, a point-push operator, and the complete handle relations have "
        "been accounted for. The calculation is delicate because a nonzero number before the full quotient has "
        "no distinguishing force: the beta, psi, and three-handle relations may kill the class."
    ),
    theorem(
        "Main claim (candidate proof)",
        "Subject to the published handle-decomposition theorems cited in Section 3 and the candidate-specific "
        "geometric identifications established in Sections 4–7, X(41,189,73) is a homotopy four-sphere and is "
        "not diffeomorphic to the standard S^4. Consequently it would be a counterexample to the smooth "
        "four-dimensional Poincare conjecture."
    ),
    para(
        "The qualifier 'candidate proof' is substantive. The complete natural-language argument and all finite "
        "certificates are public, but the Lean theorem is conditional on explicit geometry records. A clean "
        "axiom report does not discharge those records, because theorem parameters are not Lean axioms. "
        "Section 9 states this boundary precisely."
    ),
    subsection("1.1. Structure of the argument"),
    bullets([
        "Construct a balanced one-handle class v_T and a divided cubic functional D_3.",
        "Compute D_3(v_T)=-59072 and show that v_T has final quantum degree 494.",
        "Prove that D_3 kills every beta/psi relation and the six relations supplied by three chosen embedded spheres.",
        "Use the complete MWW quotient and four-handle isomorphism to obtain a nonzero degree-494 class of X(41,189,73).",
        "Use the standard-sphere calculation, which is concentrated in quantum degree zero, to exclude a diffeomorphism with S^4.",
        "Use the Cappell-Shaneson determinant criterion and standard topological arguments to identify the candidate as a homotopy sphere."
    ]),

    section("2. The Cappell–Shaneson candidate"),
    para("The matrix used throughout is"),
    equation("A = [[0, 269, 1240], [0, 41, 189], [1, 0, 32]].", "2.1"),
    para(
        "Exact integer arithmetic gives det(A)=1 and det(A-I)=1. The first identity places A in SL(3,Z); "
        "the second is the standard Cappell-Shaneson condition ensuring that the associated mapping-torus "
        "surgery has the homology and fundamental-group behavior required of a homotopy sphere. The Lean "
        "files T73Finite.lean and T73CSAlgebra.lean verify the matrix arithmetic and explicit inverses for "
        "the maps A-I and Lambda^2 A-I."
    ),
    para(
        "The topological promotion is logically separate from the smooth obstruction. Van Kampen supplies "
        "simple connectivity from the surjectivity of A-I. The Wang and Mayer-Vietoris sequences, together "
        "with Poincare duality, give the integral homology profile of S^4. Hurewicz and Whitehead then promote "
        "the result to a homotopy equivalence. In the Lean development these implications are represented by "
        "three corresponding fields of CSTopologyData; the candidate-specific applications are not presently "
        "proved inside the kernel."
    ),

    section("3. Skein-lasagna modules and external results"),
    para(
        "Let S^2_0(X;Q) denote the N=2 skein-lasagna module over Q, with its homological and quantum gradings. "
        "The proof uses the handle presentation of Manolescu, Walker, and Wedrich (MWW): a direct sum of raw "
        "cabled state spaces is quotiented first by beta/psi relations and then by the relations associated to "
        "three-handle attaching spheres. A four-handle induces a grading-preserving isomorphism."
    ),
    table(
        ["Input", "Role in this paper", "Status"],
        [
            ["MWW complete handle quotient", "Identifies the full beta/psi and sphere relation space; supplies the universal quotient.", "Published theorem; applicability must be checked."],
            ["Horvat-Jablonowski replacement", "Permits a pairwise-disjoint unimodular sphere basis to replace the historical three-handle system.", "Published/preprint theorem; candidate hypotheses checked by finite geometry."],
            ["MWW four-handle map", "Carries the surviving degree-494 class to the closed four-manifold.", "Published theorem."],
            ["MWW S^4 computation", "S^2_0(S^4;Q) is concentrated in bidegree (0,0).", "Published computation; rational scope retained explicitly."],
            ["BPW/BHPW trace functoriality", "Carries the coefficient Hattori class to the endpoint shadow and makes the required ordinary cobordism maps strict.", "Published categorical input."],
            ["Cappell-Shaneson topology", "Turns the determinant conditions into the homotopy-sphere conclusion.", "Classical topology plus the cited modern restatement."],
        ],
        [1.5, 3.35, 1.65],
    ),
    para(
        "No appeal is made to a conjectural functoriality for knotted webs or singular foams embedded in four-space. "
        "The geometric consumers in the argument are ordinary framed oriented surface cobordisms between tangles "
        "or links, together with the published horizontal- and vertical-trace constructions."
    ),

    section("4. The balanced class and the divided cubic detector"),
    subsection("4.1. Raw state and Hattori class"),
    para(
        "For a finite cabled state s=(alpha,r), let C_s be the corresponding raw MWW summand. At the selected "
        "state s_0, the active coefficient boundary has one negative and one positive cable on each of m_2 "
        "and r_xy. Thus alpha=0, r is the sum of the two corresponding basis vectors, and |r|=2."
    ),
    equation("C(s) = S²₀(W₁; K(r−α⁻, r+α⁺) ∪ L, ηʳ; Q){−(2|r|+|α|)}.", "4.1"),
    para(
        "The actual paired-annulus cut yields an action-compatible coefficient equivalence. Here B_act is the "
        "framed automorphism of the 88-endpoint object induced by the paired-annulus motion. The two transported "
        "boundary objects are T_1=B_act^{-1}U and T_0=B_act^{-1}WU, where U is the physical cup and W is the "
        "point-push braid. The selected diagonal Hattori class is"
    ),
    equation("v_T = η_{R}[T_1] = [ H^{-1}(Id_{U} ⊗ X^{⊗ 227}) ].", "4.2"),
    para(
        "The 227 symbols X label separate Frobenius-algebra factors; they are not the vanishing product X^227. "
        "The coefficient trace and 227 counits carry this class to the endpoint vector u=e_0-e_5. The discarded "
        "alternative ξ=η_{R}[T_0] − s_{inv}η_{R}[T_1] is not used: applying the same outer point-push difference "
        "to ξ introduces a second factor and first contributes in order six."
    ),
    subsection("4.2. Quantum trace and divided functional"),
    para(
        "Let zeta be the quantum-trace parameter and complete at zeta=1+h. The point-push word lies in the third "
        "term of the relevant filtration, so rho_h(W)−I is divisible by h^3 on the full endpoint module. With a "
        "normalized cap Ĉ(h), division by h^3 followed by reduction modulo h defines an ordinary linear "
        "functional D_3 on coefficient HH_0. The map Theta_h is the completed quantum-Hochschild shadow from "
        "the actual coefficient module to the endpoint module."
    ),
    equation("D₃ = [h³] Ĉ(h) (ρ_h(W)−I) P₈₆ Θ_h.", "4.3"),
    para(
        "The projector P_86 isolates the one-cup top cell. It is applied only after the genuine quantum-Hochschild "
        "shadow; it is not asserted to be a physical foam. The construction is lift-independent because changing "
        "a lift by h y changes the divided value only by a multiple of h."
    ),
    subsection("4.3. Exact point-push calculation"),
    para("Put t=zeta^{-2} and epsilon=t-1. Exact integer arithmetic gives"),
    equation("ℓ(ρ(W)−I)u = 7384 ε³ − 660412 ε⁴ + 34814626 ε⁵ − 1365512573 ε⁶.", "4.4"),
    para(
        "The coefficients in degrees zero, one, and two vanish. Since epsilon=-2h+3h^2-4h^3+O(h^4), the cubic "
        "coefficient is"
    ),
    equation("D₃(v_T) = (−2)³ · 7384 = −59072 ≠ 0.", "4.5"),
    para(
        "The public recomputation starts from 252 primitive crossing rows and a chronology record. It reconstructs "
        "an 11,340-letter B44 word, then its 45,360-letter two-cable word, and performs the truncated Burau action. "
        "Neither 7384 nor -59072 appears in the input JSON. The value belongs to the registered "
        "oriented point-push chronology. A different, non-chronological ordering of the same crossing rows gives "
        "-58976 and moves the first anomaly from degree three to degree two. The legality and naturality arguments "
        "must therefore establish invariance under every permitted change of presentation; the finite number alone "
        "is not yet a manifold invariant."
    ),
    subsection("4.4. Grading"),
    para("The four grading contributions are exact:"),
    equation("-44 + 227 + 315 - 4 = 494.", "4.6"),
    para("Hence the proposed surviving class lies in bidegree (0,494)."),

    section("5. Descent through the one- and two-handle relations"),
    subsection("5.1. One-cup firewall"),
    para(
        "The endpoint Bar-Natan/Temperley-Lieb category is filtered by through degree: F_k is the span of diagrams "
        "with through degree at most k, and composition cannot raise this filtration. The physical cup U has "
        "through degree 86. Every undotted balanced-pair creation for an active gate-crossing owner contains at "
        "least two cups, so its full action-closed ideal lies in F_84. "
        "The rational projector P_86 therefore kills that ideal while preserving the selected vector."
    ),
    para(
        "The zero-gate owner r_zx is handled separately. After the artificial meridian spectator is deleted, its "
        "balanced pair is a split zero-framed unknot. For the standard rank-two Frobenius algebra A=Q[X]/(X^2),"
    ),
    equation("(ε ⊗ ε) ψ⁽⁰⁾ = 0,      (ε ⊗ ε) ψ⁽¹⁾ = Id.", "5.1"),
    subsection("5.2. All finite cable states"),
    para(
        "At the base state the constant balanced-pair quotient is trivial and every pure generator acts as I+O(h). "
        "At higher multiplicities, physical copies are put into a fixed order by stable positive shuffles, namely "
        "positive braids that preserve the order already assigned to existing copies. After the actual W2 core "
        "disks are attached, the functional is averaged over the finite permutation orbit; this is the Reynolds "
        "average used here. The normalization ratios for successive orbit sizes telescope, so the intermediate "
        "factors cancel. Distinct-owner squares commute because their supports are disjoint. Consequently, for "
        "every finite state and every beta/psi edge,"
    ),
    equation("Λₜ⁽³⁾ Ψₑ⁰ = 0,      Λₜ⁽³⁾ Ψₑ¹ = Λₛ⁽³⁾.", "5.2"),
    para(
        "Equations (5.1)-(5.2) imply that the divided functional annihilates the complete one- and two-handle "
        "relation subspace. This is an all-state statement at order h^3, not a test on the selected vector alone."
    ),

    section("6. A chosen three-sphere basis"),
    subsection("6.1. Finite geometry"),
    para(
        "The historical attaching-sphere movies are not recovered and are not used. Instead, the proof constructs "
        "three new chosen spheres, denoted TH1, TH2, and THXY, in the same post-two-handle boundary. Their frozen "
        "certificates and hostile re-audits are included in the public evidence directory."
    ),
    table(
        ["Sphere", "Construction counts", "Frozen SHA-256 prefix"],
        [
            ["TH1", "350,176 leaves; 350,175 bands; one root cap", "EE620E6B085A5F9..."],
            ["TH2", "229,198 leaves; 229,197 bands; one root cap", "4D1B627C0343A1C..."],
            ["THXY", "11,115 material disks; 11,114 split bands; one root cap", "EABF67C0D1AE309..."],
        ],
        [0.8, 4.0, 1.7],
    ),
    para(
        "The certificates record complete ancestry hashes, embedded product-normal bands, root-cap consumption, "
        "and Euler-characteristic ledgers (χ = 2, i.e. genus zero). The three movie-time (height) intervals are disjoint: THXY lies in [2,13/2], TH1 in [8,9], "
        "and TH2 in [10,11], while the old one-cup/core block lies in [-4,-3]."
    ),
    subsection("6.2. Homology basis and replacement"),
    para("Their spherical homology coordinates are"),
    equation("v₁=(−1311,8608,−1),   v₂=(−189,1241,0),   v₃=(41,−269,1).", "6.1"),
    equation("det[v₁  v₂  v₃] = 1.", "6.2"),
    para(
        "The Horvat–Jabłonowski basis theorem is then invoked to replace the historical three-handle attaching "
        "system by this pairwise-disjoint unimodular basis, up to permutation and three-handle slides. The external "
        "link L is empty in the present application. This replacement theorem is an external published/preprint "
        "input; the embeddedness, disjointness, and determinant-one hypotheses are candidate-specific finite claims."
    ),
    subsection("6.3. The six sphere relations"),
    para(
        "For each chosen sphere j, every noninvertible critical point lies in a new material factor. The old one-cup "
        "block is transported by identity cylinders. After the fixed actual-to-canonical coordinate maps, the "
        "constant endpoint maps have the form"
    ),
    equation("F(j,0)⁽⁰⁾ = Q(t_j,0)⁻¹ (Id_old ⊗ U_j) Q(s,0),", "6.3"),
    equation("F(j,1)⁽⁰⁾ = Q(t_j,0)⁻¹ (Id_old ⊗ D_j) Q(s,0),", "6.4"),
    equation("D_j=X^{⊗ b_j};   U_j=∑_{a=0}^{b_j−1} X^{⊗ a} ⊗ 1 ⊗ X^{⊗(b_j−1−a)}.", "6.5"),
    para("The W2 core disks give E_j as the b_j-fold tensor power of the counit epsilon, whence E_j(U_j)=0 and E_j(D_j)=1. Write Ψ^quot(j,n) for the map induced after the beta/psi quotient. Therefore"),
    equation("ℓ₁ Ψ^quot(j,0)=0,      ℓ₁(Ψ^quot(j,1)−Id)=0,      j=1,2,3.", "6.6"),
    para(
        "These are identities of linear maps on the whole source under the stated naturality square. The stored "
        "scalar pairs 0/0, 0/0, 0/0 are consequences, not substitutes for (6.6). Positive transport corrections "
        "start at h^4 after multiplication by the old h^3 anomaly and therefore do not alter the divided cubic. "
        "This order estimate concerns endpoint transport. It does not by itself settle a change in the insertion "
        "position of a collar word; that separate geometric binding is listed in Section 11."
    ),

    section("7. Passage to the closed manifold"),
    para(
        "We use three levels of notation. D₃ denotes the numerical value of the detector on the selected class; "
        "Λ₃ denotes the compatible functional on the full direct sum of raw states; Λ₃^quot denotes its "
        "descent to the complete quotient. The Lean fields ell0, ell1, and ell2 are the corresponding functionals "
        "at the successive quotient stages."
    ),
    para(
        "Let C be the direct sum of all finite raw cable states, R_2h the complete beta/psi relation subspace, "
        "and R_3h the six-relation subspace from the chosen sphere basis. Sections 5 and 6 assert"
    ),
    equation("R₂h ⊂ ker Λ₃,      R₃h ⊂ ker Λ₃.", "7.1"),
    para("Write v_raw(s_0) for the raw-state representative of v_T in the selected state s_0."),
    para("The quotient universal property produces"),
    equation("Λ₃^quot : C/(R₂h+R₃h) → Q,      Λ₃=Λ₃^quot π.", "7.2"),
    equation("Λ₃^quot(π(v_raw(s_0))) = −59072 ≠ 0.", "7.3"),
    para(
        "Thus the selected degree-494 class survives the complete handle quotient. The published four-handle "
        "isomorphism preserves its bidegree and nonvanishing. For the standard four-sphere, the same invariant is "
        "concentrated in bidegree (0,0); in particular S^2_{0,494}(S^4;Q)=0. Graded diffeomorphism invariance would "
        "carry a degree-494 class of the candidate to one of S^4, which is impossible. The standard control itself "
        "has W_std=I and thus checks only a zero-to-zero map; it cannot independently diagnose a failure of the "
        "candidate descent."
    ),
    theorem(
        "Smooth obstruction",
        "If the actual-geometry identifications, complete MWW quotient, chosen-sphere replacement, four-handle "
        "comparison, and rational S^4 control have the scopes stated above, then X(41,189,73) is not "
        "diffeomorphic to S^4."
    ),

    section("8. Evidence and independent replay"),
    para(
        "The public repository separates proof prose, exact finite inputs, replay programs, and Lean sources. "
        "The initial public release mistakenly retained several load-bearing geometry objects only by local path "
        "and SHA-256. Commit 9797850472602f311b6957244048044d70b1adb4 corrects that omission."
    ),
    table(
        ["Layer", "Public artifact", "What can be independently checked"],
        [
            ["Point-push cubic", "T73_DELTA3_PUBLIC_INPUT.json; recompute_t73_delta3.py", "Rebuild B44/B88; exact epsilon series; -59072."],
            ["G1 global descending", "t73_reduced_billiard.pd.json; verify_t73_global_descending.py", "All 2,126,291 crossings, components, basepoints, height order."],
            ["E1 coefficient cut", "ACTUAL_PD_CABLE_UNIT_CERT.json", "Exact passage and disk certificate; frozen input identity."],
            ["E8 chosen spheres", "TH1/TH2/THXY JSON plus hostile re-audits", "Actual frozen certificate contents, hashes, counts, determinant and scalar fields."],
            ["Lean kernel", "Smooth4PC/*.lean; T73Audit.lean", "Arithmetic, abstract trace/quotient algebra, final conditional implication, axiom reports."],
        ],
        [1.15, 2.7, 2.65],
    ),
    para(
        "The public geometry verifier checks 29 immutable files and recomputes the global-descending certificate "
        "from the full PD. It does not regenerate the three chosen-sphere certificates from their much larger "
        "upstream construction trees. Those JSON files are now available for independent inspection, but their "
        "geometric semantics remain part of the human proof rather than the Lean kernel."
    ),

    section("9. What the Lean development proves"),
    subsection("9.1. Actual mathlib use"),
    para(
        "The project pins mathlib at revision 520045ab14e26149ee970e2e617ca04b09bde5d6. It imports rational "
        "linear algebra, quotient modules, finite sums, the fundamental-groupoid simply-connectedness predicate, "
        "and singular homology. The project does use mathlib; the substantive question is whether the "
        "candidate-specific geometry has been constructed as Lean data."
    ),
    subsection("9.2. Kernel-checked content"),
    bullets([
        "The determinant and lattice calculations for A, A-I, and Lambda^2 A-I.",
        "The exact arithmetic identities computedCubic=-59072 and computedDegree=494 from frozen integer inputs.",
        "Generic coefficient-HH_0 quotient and descended-row naturality lemmas.",
        "The positive split-tree Frobenius identities: the undotted row is zero and the dotted row is the source row.",
        "Generic sphere-relation kernel inclusion and the implication from a nonzero descended row to a nonzero quotient class.",
        "The final logical implication from an ExternalGeometry and Cappell-Shaneson topology package to 'homotopy sphere and not diffeomorphic to S^4'."
    ]),
    subsection("9.3. Explicit formalization boundary"),
    para("The final theorem has the schematic type"),
    code([
        "conditionalCounterexample",
        "  (geom : ExternalGeometry u)",
        "  (cs   : CSExternalGeometry u) :",
        "  IsHomotopySphere candidate /\\ not Diffeomorphic candidate S4",
    ]),
    para(
        "No tracked file constructs geom or cs. ExternalGeometry contains W0-W3, the selected class and functional, "
        "the equation ell0(x0)=-59072, quotient maps q01 and q12 with compatibility equations, a transport "
        "equivalence, the four-handle equivalence, vanishing of the standard-sphere degree, and graded "
        "diffeomorphism invariance. Related interfaces require actual sphere embeddings, class coordinates, six "
        "actual sphere maps, pairwise disjointness, HJ replacement, the MWW coequalizer, and the four-handle binding."
    ),
    para(
        "Likewise, CSTopologyData uses genuine mathlib notions of simply connectedness and singular homology but "
        "takes the candidate-specific van Kampen, Wang/Mayer-Vietoris/Poincare, and Hurewicz/Whitehead implications "
        "as fields. The local file named Unconditional.lean contains documentation only and expressly says that "
        "the current result remains conditional."
    ),
    subsection("9.4. Why the axiom audit is necessary but insufficient"),
    para(
        "A clean '#print axioms' report establishes that the proofs do not use sorryAx or undeclared project axioms. "
        "The fresh replay produces 38 reports, all contained in {propext, Classical.choice, Quot.sound}. This is "
        "useful kernel evidence. It does not show that ExternalGeometry is inhabited: an explicit theorem parameter "
        "is an assumption in the theorem type, not an axiom reported by '#print axioms'."
    ),

    section("10. Reproduction protocol"),
    para("At repository commit 9797850472602f311b6957244048044d70b1adb4, a reader may run:"),
    code([
        "git clone https://github.com/toffee-desuwa/smooth4pc-t73-lean.git",
        "cd smooth4pc-t73-lean",
        "lake update",
        "git diff --exit-code -- lake-manifest.json",
        "lake exe cache get",
        "python -B tests/test_t73_minimal_formalization.py -v",
        "lake lean T73Audit.lean",
        "python -I -B scripts/recompute_t73_delta3.py --check",
        "python -I -B scripts/verify_public_geometry_evidence.py",
    ]),
    para(
        "The expected outputs are: two Python formalization tests pass; T73Audit exits zero; 38 axiom reports "
        "contain only the three foundational axioms listed above; the point-push replay prints "
        "DELTA3_ETA_T1=-59072; and the geometry replay prints PD_CROSSINGS=2126291, "
        "GLOBAL_DESCENDING=PASS, and VERIFY=PASS."
    ),

    section("11. Independent-review priorities"),
    para(
        "The finite arithmetic is the least uncertain part of the argument. A serious review should concentrate "
        "on the interfaces that connect finite certificates and published theorems to the actual smooth "
        "four-manifold. The following questions are decisive."
    ),
    bullets([
        "Does the actual paired-annulus cut give the claimed two-sided Hattori coefficient, including framing and all left/right actions?",
        "Does the quantum vertical-to-horizontal trace apply with the stated grading and strict naturality to every ordinary cobordism used here?",
        "Does the actual-Gompf-to-DIAGRAM bridge preserve the complete labelled framed link, rather than only relator words or unframed projections?",
        "Do the TH1, TH2, and THXY certificates define three embedded, framed, pairwise-disjoint spheres in one actual W2, with the displayed homology coordinates?",
        "Does the Horvat–Jabłonowski replacement theorem apply with exactly these hypotheses and preserve the upper cobordism relative to W2?",
        "Are the six actual sphere maps identified with the canonical new-factor split maps on the whole source, not merely on the selected vector?",
        "Does the complete MWW quotient contain no additional relation that can kill the class?",
        "Are the four-handle and rational S^4 comparisons grading-preserving in precisely the form consumed by the proof?",
        "The value D_3=-59072 is computed for the registered oriented point-push presentation and serialization. Reordering the same crossing data can give -58976 and move the first anomaly from degree three to degree two. Which equivalence theorem proves invariance under every legal change of presentation and ordering?",
        "The standard S^4 control has W_std=I and hence sends zero to zero. Because that control cannot expose a failure of descent, is the candidate descent proved independently rather than inferred from agreement with the control?",
        "Is the insertion position of each sphere movie relative to the detector fixed by the geometric binding? A collar ambiguity w -> w p with p in the third term Γ₃ of the lower central series of the pure braid group can be visible at order h^3; the observed value is divisible by 16, so the insertion-point issue cannot be dismissed by the present congruence alone.",
    ]),
    para(
        "A failure at any of these points breaks the claimed smooth obstruction even though all current Lean files "
        "continue to compile. Conversely, verifying these identifications independently would convert the present "
        "conditional kernel into strong support for the full mathematical argument."
    ),

    section("12. Conclusion"),
    para(
        "The proposed counterexample is detected by a sharply localized discrepancy: a degree-494 class whose "
        "divided cubic value is -59072. The calculation is exact, the complete finite evidence is now public, and "
        "the abstract descent implication has been checked in Lean. What remains outside the kernel is not a vague "
        "appeal to intuition but a finite list of named geometric and published-theorem interfaces. This separation "
        "is the principal value of the present package: it makes clear both why the candidate would falsify the "
        "conjecture and exactly where an independent mathematical review must succeed or find an error."
    ),

    section("Appendix A. Formal interface inventory"),
    table(
        ["Interface field", "Mathematical content", "Current support"],
        [
            ["ell0_x0", "The selected actual class evaluates to -59072.", "Exact Python replay plus human binding; explicit Lean field."],
            ["q01, q12", "Complete two successive quotient maps.", "Abstract quotient lemmas in Lean; actual identification external."],
            ["ell1_comp_q01, ell2_comp_q12", "Detector descends through both relation spaces.", "Generic Lean implication; candidate naturality human-checked."],
            ["transport, fourIso", "MWW transport and four-handle comparison.", "Published input; explicit Lean field."],
            ["s4DegreeZero", "The S^4 group in degree 494 vanishes.", "Published computation; explicit Lean field."],
            ["diffeomorphismEquiv", "Diffeomorphism induces a graded equivalence.", "Invariant functoriality; explicit Lean field."],
            ["sphere bindings", "Actual embedded spheres induce the six claimed maps.", "Public JSON/prose evidence; not instantiated in Lean."],
            ["CSTopologyData", "Simple connectivity, homology-sphere profile, and homotopy promotion.", "Mathlib definitions; candidate implications supplied as fields."],
        ],
        [1.4, 2.55, 2.55],
    ),

    section("Appendix B. Core immutable identities"),
    code([
        "T73 delta input       04507506DC577384BC4C04765CCB212C1481DC71810C6AB3232F1AB690F16909",
        "B44 derived word      7C2D2F792C2672221A76CAF08A71F560AF0CB7654B4D537C00FAD00B16EFA187",
        "B88 derived word      C13D0A3BA9B05F41A6D2C5B4AB12DDECDABEFC1883835891AD3BD2B8955B5FFB",
        "Actual cable unit     7F3D3618D6A790A9B60EE8085B647AC2AB742E1BC9C15841F1BEF015034217B5",
        "Global PD             E6912A64457557469E5C691B4D57ABDBBF4C45ADB05492777C574223D0C06F8A",
        "TH1 sphere            EE620E6B085A5F9E1C73CFDD1AD04FC0682CEC74DA3DBF8AFE70DD19C038E3A0",
        "TH2 sphere            4D1B627C0343A1C464319704EAFADCA127902A2E4E90CF1C283004359B7ADC24",
        "THXY sphere           EABF67C0D1AE309F0281297710A283B8ED5A11D88C8E810837932313F4C53227",
    ]),

    section("References"),
    {"kind": "references", "items": [
        "[1] D. Aitchison and J. H. Rubinstein, Fibered knots and involutions on homotopy spheres, in Four-Manifold Theory, Contemporary Mathematics 35, 1984.",
        "[2] A. Beliakova, M. Hogancamp, K. Putyra, and S. Wehrli, On the functoriality of sl(2) tangle homology, Algebr. Geom. Topol. 23 (2023), 1303–1361; arXiv:1903.12194.",
        "[3] A. Beliakova, K. Putyra, and S. Wehrli, Quantum link homology via trace functor I, arXiv:1605.03523.",
        "[4] A. Hatcher, Algebraic Topology, Cambridge University Press, 2002; official online chapters at pi.math.cornell.edu/~hatcher/AT/.",
        "[5] E. Horvat and M. Jabłonowski, On 4-dimensional 3-handle attachments, arXiv:2510.20282.",
        "[6] M. Iwaki, Infinite families of standard Cappell-Shaneson spheres, Topology and its Applications 366C (2025), 109293; arXiv:2404.05096.",
        "[7] F. Laudenbach and V. Poenaru, A note on 4-dimensional handlebodies, Bulletin de la Societe Mathematique de France 100 (1972), 337–344.",
        "[8] C. Manolescu and I. Neithalath, Skein lasagna modules for 2-handlebodies, Journal fur die reine und angewandte Mathematik; arXiv:2009.08520.",
        "[9] C. Manolescu, K. Walker, and P. Wedrich, Skein lasagna modules and handle decompositions, Advances in Mathematics (2023), 109071; arXiv:2206.04616.",
        "[10] S. Morrison, K. Walker, and P. Wedrich, Invariants of 4-manifolds from Khovanov-Rozansky link homology, Geometry & Topology 26 (2022), 3367–3420; arXiv:1907.12194.",
    ]},
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, name: str = "Times New Roman", size: float | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Courier New")
    fonts.set(qn("w:hAnsi"), "Courier New")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A4B8C")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    props.extend([fonts, color, underline, size])
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after in (
        ("Heading 1", 13.0, 14, 6),
        ("Heading 2", 11.0, 10, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("T73 SKEIN-LASAGNA CANDIDATE OBSTRUCTION")
    set_run_font(run, size=8)
    run.font.color.rgb = RGBColor(90, 90, 90)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Public review copy")
    set_run_font(run, size=8)
    run.font.color.rgb = RGBColor(90, 90, 90)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(TITLE)
    set_run_font(r, size=17, bold=True)
    for item in CONTENT:
        kind = item["kind"]
        if kind == "abstract":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run("Abstract. ")
            set_run_font(r, size=9.5, bold=True)
            r = p.add_run(pretty_math(item["text"]))
            set_run_font(r, size=9.5)
        elif kind == "keywords":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run("Keywords. ")
            set_run_font(r, size=9, bold=True)
            r = p.add_run(pretty_math(item["text"]))
            set_run_font(r, size=9)
        elif kind == "section":
            if item["text"].startswith("Appendix B."):
                doc.add_page_break()
            doc.add_heading(item["text"], level=1)
        elif kind == "subsection":
            doc.add_heading(item["text"], level=2)
        elif kind == "para":
            p = doc.add_paragraph(pretty_math(item["text"]))
            p.paragraph_format.keep_together = False
        elif kind == "equation":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(pretty_math(item["text"]) + (f"    ({item['number']})" if item["number"] else ""))
            set_run_font(r, name="Cambria Math", size=10.2, italic=True)
        elif kind == "theorem":
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
            cell = t.cell(0, 0)
            cell.width = Inches(6.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "F1F1F1")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(item["label"] + ". ")
            set_run_font(r, size=10.2, bold=True)
            r = p.add_run(pretty_math(item["text"]))
            set_run_font(r, size=10.2, italic=True)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == "bullets":
            for value in item["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.15)
                p.paragraph_format.space_after = Pt(3)
                p.add_run(pretty_math(value))
        elif kind == "code":
            for line in item["lines"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.space_after = Pt(0)
                url = "https://github.com/toffee-desuwa/smooth4pc-t73-lean.git"
                if url in line:
                    prefix, suffix = line.split(url, 1)
                    set_run_font(p.add_run(prefix), name="Courier New", size=8.2)
                    add_hyperlink(p, url, url)
                    if suffix:
                        set_run_font(p.add_run(suffix), name="Courier New", size=8.2)
                else:
                    set_run_font(p.add_run(line), name="Courier New", size=8.2)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(item["headers"]))
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
            for idx, (header_text, width) in enumerate(zip(item["headers"], item["widths"])):
                cell = t.rows[0].cells[idx]
                cell.width = Inches(width)
                set_cell_shading(cell, "E8E8E8")
                r = cell.paragraphs[0].add_run(header_text)
                set_run_font(r, size=8.5, bold=True)
            for row in item["rows"]:
                cells = t.add_row().cells
                for idx, (value, width) in enumerate(zip(row, item["widths"])):
                    cells[idx].width = Inches(width)
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    p = cells[idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    set_run_font(p.add_run(pretty_math(value)), size=8.2)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == "references":
            for value in item["items"]:
                p = doc.add_paragraph(pretty_math(value))
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(3)
                for r in p.runs:
                    set_run_font(r, size=8.8)

    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "Candidate Smooth4PC falsification proof and formalization boundary"
    doc.core_properties.keywords = "Smooth4PC, Cappell-Shaneson, skein lasagna, Lean"
    doc.save(DOCX)


def register_pdf_fonts() -> None:
    fonts = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("TNR", str(fonts / "times.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-Bold", str(fonts / "timesbd.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-Italic", str(fonts / "timesi.ttf")))
    pdfmetrics.registerFont(TTFont("TNR-BoldItalic", str(fonts / "timesbi.ttf")))
    pdfmetrics.registerFont(TTFont("CambriaMath", str(fonts / "cambria.ttc"), subfontIndex=1))
    pdfmetrics.registerFontFamily(
        "TNR", normal="TNR", bold="TNR-Bold", italic="TNR-Italic", boldItalic="TNR-BoldItalic"
    )


class PaperDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str):
        super().__init__(
            filename,
            pagesize=letter,
            leftMargin=0.85 * inch,
            rightMargin=0.85 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.72 * inch,
            title=TITLE,
            author=AUTHOR,
        )
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body")
        self.addPageTemplates(PageTemplate(id="paper", frames=frame, onPage=self.page_furniture))

    @staticmethod
    def page_furniture(canvas, doc):
        canvas.saveState()
        canvas.setFont("TNR", 7.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        if doc.page > 1:
            canvas.drawCentredString(letter[0] / 2, letter[1] - 0.38 * inch,
                                     "T73 SKEIN-LASAGNA CANDIDATE OBSTRUCTION")
        canvas.drawCentredString(letter[0] / 2, 0.36 * inch, str(doc.page))
        canvas.restoreState()


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(pdf_markup(text).replace("\n", "<br/>"), style)


def build_pdf() -> None:
    register_pdf_fonts()
    base = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "PaperBody", parent=base["BodyText"], fontName="TNR", fontSize=9.6,
            leading=12.4, alignment=TA_JUSTIFY, spaceAfter=6,
        ),
        "section": ParagraphStyle(
            "PaperSection", parent=base["Heading1"], fontName="TNR-Bold", fontSize=12.3,
            leading=14.8, textColor=colors.black, spaceBefore=12, spaceAfter=5,
            keepWithNext=True,
        ),
        "subsection": ParagraphStyle(
            "PaperSubsection", parent=base["Heading2"], fontName="TNR-Bold", fontSize=10.4,
            leading=12.5, textColor=colors.black, spaceBefore=9, spaceAfter=4,
            keepWithNext=True,
        ),
        "abstract": ParagraphStyle(
            "Abstract", parent=base["BodyText"], fontName="TNR", fontSize=8.9,
            leading=11.3, alignment=TA_JUSTIFY, leftIndent=0.32*inch,
            rightIndent=0.32*inch, spaceAfter=7,
        ),
        "equation": ParagraphStyle(
            "Equation", parent=base["BodyText"], fontName="CambriaMath", fontSize=9.7,
            leading=12, alignment=TA_CENTER, spaceBefore=4, spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "PaperBullet", parent=base["BodyText"], fontName="TNR", fontSize=9.3,
            leading=11.8, alignment=TA_LEFT, leftIndent=0.28*inch,
            firstLineIndent=-0.14*inch, bulletIndent=0.05*inch, spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "PaperCode", parent=base["Code"], fontName="Courier", fontSize=7.3,
            leading=9.0, leftIndent=0.25*inch, rightIndent=0.1*inch,
            backColor=colors.HexColor("#F5F5F5"), borderPadding=4, spaceAfter=0,
        ),
        "reference": ParagraphStyle(
            "Reference", parent=base["BodyText"], fontName="TNR", fontSize=7.6,
            leading=8.8, alignment=TA_LEFT, leftIndent=0.22*inch,
            firstLineIndent=-0.22*inch, spaceAfter=2,
        ),
        "table": ParagraphStyle(
            "TableText", parent=base["BodyText"], fontName="TNR", fontSize=7.6,
            leading=9.2, alignment=TA_LEFT,
        ),
        "tablehead": ParagraphStyle(
            "TableHead", parent=base["BodyText"], fontName="TNR-Bold", fontSize=7.8,
            leading=9.4, alignment=TA_LEFT,
        ),
    }

    story = []
    story.append(Spacer(1, 0.22 * inch))
    story.append(Paragraph(escape(TITLE), ParagraphStyle(
        "Title", fontName="TNR-Bold", fontSize=17.5, leading=21,
        alignment=TA_CENTER, spaceAfter=14,
    )))

    for item in CONTENT:
        kind = item["kind"]
        if kind == "abstract":
            story.append(Paragraph("<b>Abstract.</b> " + pdf_markup(item["text"]), styles["abstract"]))
        elif kind == "keywords":
            story.append(Paragraph("<b>Keywords.</b> " + pdf_markup(item["text"]), styles["abstract"]))
        elif kind == "section":
            if item["text"].startswith("Appendix B."):
                story.append(PageBreak())
            story.append(Paragraph(escape(item["text"]), styles["section"]))
        elif kind == "subsection":
            story.append(Paragraph(escape(item["text"]), styles["subsection"]))
        elif kind == "para":
            story.append(pdf_paragraph(item["text"], styles["body"]))
        elif kind == "equation":
            value = item["text"] + (f"     ({item['number']})" if item["number"] else "")
            story.append(pdf_paragraph(value, styles["equation"]))
        elif kind == "theorem":
            inner = Paragraph("<b>" + escape(item["label"]) + ".</b> <i>" + pdf_markup(item["text"]) + "</i>", styles["body"])
            box = Table([[inner]], colWidths=[6.25*inch], hAlign="CENTER")
            box.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F1F1F1")),
                ("BOX", (0,0), (-1,-1), 0.5, colors.HexColor("#888888")),
                ("LEFTPADDING", (0,0), (-1,-1), 8),
                ("RIGHTPADDING", (0,0), (-1,-1), 8),
                ("TOPPADDING", (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 1),
            ]))
            story.extend([KeepTogether(box), Spacer(1, 6)])
        elif kind == "bullets":
            for value in item["items"]:
                story.append(Paragraph("• " + pdf_markup(value), styles["bullet"]))
        elif kind == "code":
            joined = "<br/>".join(code_markup(line) if line else "&nbsp;" for line in item["lines"])
            inner = Paragraph(joined, styles["code"])
            block = Table([[inner]], colWidths=[6.15*inch], hAlign="CENTER")
            block.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F5F5F5")),
                ("LEFTPADDING", (0,0), (-1,-1), 5),
                ("RIGHTPADDING", (0,0), (-1,-1), 5),
                ("TOPPADDING", (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ]))
            story.append(KeepTogether(block))
            story.append(Spacer(1, 6))
        elif kind == "table":
            rows = [[Paragraph(escape(h), styles["tablehead"]) for h in item["headers"]]]
            rows += [[Paragraph(pdf_markup(v), styles["table"]) for v in row] for row in item["rows"]]
            tbl = Table(rows, colWidths=[w*inch for w in item["widths"]], repeatRows=1, hAlign="CENTER")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8E8E8")),
                ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#777777")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING", (0,0), (-1,-1), 4),
                ("RIGHTPADDING", (0,0), (-1,-1), 4),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.extend([tbl, Spacer(1, 7)])
        elif kind == "references":
            for value in item["items"]:
                story.append(pdf_paragraph(value, styles["reference"]))

    PaperDocTemplate(str(PDF)).build(story)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(f"DOCX={DOCX}")
    print(f"PDF={PDF}")
